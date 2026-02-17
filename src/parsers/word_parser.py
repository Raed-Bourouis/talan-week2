"""Word document parser."""
import logging
from typing import Dict, Any, List
from docx import Document
from docx.table import Table

from . import BaseParser

logger = logging.getLogger(__name__)


class WordParser(BaseParser):
    """Parser for Word documents."""
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """Parse a Word document and extract structured data."""
        try:
            text = self.extract_text(file_path)
            paragraphs = self._extract_paragraphs(file_path)
            tables = self._extract_tables(file_path)
            
            return {
                "text": text,
                "paragraphs": paragraphs,
                "tables": tables,
                "file_type": "word"
            }
        except Exception as e:
            logger.error(f"Error parsing Word document {file_path}: {e}")
            raise
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from a Word document."""
        try:
            doc = Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join([cell.text for cell in row.cells])
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting text from Word document {file_path}: {e}")
            raise
    
    def _extract_paragraphs(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract paragraphs with formatting information."""
        try:
            doc = Document(file_path)
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        "text": para.text,
                        "style": para.style.name if para.style else "Normal"
                    })
            
            return paragraphs
        except Exception as e:
            logger.warning(f"Error extracting paragraphs from Word document {file_path}: {e}")
            return []
    
    def _extract_tables(self, file_path: str) -> List[List[List[str]]]:
        """Extract tables from Word document."""
        try:
            doc = Document(file_path)
            tables = []
            
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            return tables
        except Exception as e:
            logger.warning(f"Error extracting tables from Word document {file_path}: {e}")
            return []
